import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from generator import PAGE_BREAK_MARKER, normalize_pdf_breaks, normalize_quick_take_text
from pdf_generator import (
    _is_page_break_line,
    _parse_role_block,
    strip_coaching_notes,
)

TEMPLATE_PATH = Path(__file__).resolve().parent / "assets" / "template.docx"
NAME_FONT = "Helvetica Neue"

# Section titles match assets/template.html (PDF) exactly
SECTION_SUMMARY = "Summary"
SECTION_SKILLS = "Skills"
SECTION_EXPERIENCE = "Experience"
SECTION_SCHOOL_PROJECTS = "School Projects"
SECTION_CREDENTIALS = "Credentials"


def _doc_style_tokens(export_mode: str = "standard") -> dict:
    if export_mode == "digital":
        return {
            "name_size": 18,
            "tagline_size": 11,
            "contact_size": 9,
            "heading_space_before": 14,
            "heading_space_after": 6,
            "body_space_after": 9,
            "bullet_space_after": 7,
            "role_space_before": 11,
            "role_summary_space_after": 10,
            "company_desc_space_after": 6,
        }
    return {
        "name_size": 17,
        "tagline_size": 10.5,
        "contact_size": 9,
        "heading_space_before": 12,
        "heading_space_after": 5,
        "body_space_after": 8,
        "bullet_space_after": 6,
        "role_space_before": 10,
        "role_summary_space_after": 9,
        "company_desc_space_after": 5,
    }


def extract_coaching_notes(text):
    """Extract only COACHING NOTE lines from text for UI display."""
    notes = []
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("COACHING NOTE:") or stripped.startswith("Alt (general)"):
            notes.append(stripped)
    return "\n".join(notes)


def _clear_body(doc):
    """Remove all paragraphs and tables from document body."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl", "sdt"):
            body.remove(child)


def _strip_html_tags(text: str) -> str:
    return re.sub(r"<[^>]+>", "", text or "").strip()


def _set_paragraph_bottom_border(paragraph, size=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_para(
    doc,
    text,
    style="Normal",
    align=None,
    bold=False,
    italic=False,
    font_size=None,
    font_name=None,
    keep_with_next=False,
    page_break_before=False,
    space_after_pt=None,
    space_before_pt=None,
):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if align is not None:
        p.alignment = align
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    if page_break_before:
        p.paragraph_format.page_break_before = True
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    if space_before_pt is not None:
        p.paragraph_format.space_before = Pt(space_before_pt)
    return p


def _add_section_heading(doc, title: str, page_break_before: bool = False, tokens: dict | None = None):
    """Match template.html h2: uppercase, 10.5pt, bottom rule."""
    tokens = tokens or _doc_style_tokens()
    p = doc.add_paragraph(style="Normal")
    if page_break_before:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(tokens["heading_space_before"])
    p.paragraph_format.space_after = Pt(tokens["heading_space_after"])
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = NAME_FONT
    _set_paragraph_bottom_border(p)
    return p


def _add_runs_with_metric_bold(paragraph, text: str, italic=False):
    from generator import find_first_metric_span

    span = find_first_metric_span(text)
    if not span:
        run = paragraph.add_run(text)
        if italic:
            run.italic = True
        return
    start, end = span
    if start > 0:
        run = paragraph.add_run(text[:start])
        if italic:
            run.italic = True
    metric_run = paragraph.add_run(text[start:end])
    metric_run.bold = True
    if italic:
        metric_run.italic = True
    if end < len(text):
        run = paragraph.add_run(text[end:])
        if italic:
            run.italic = True


def _add_list_bullet(doc, line: str, tokens: dict | None = None):
    """Bulleted line with optional bold label and first metric (Summary method bullets / role bullets)."""
    tokens = tokens or _doc_style_tokens()
    stripped = line.strip()
    if stripped.startswith(("-", "•")):
        stripped = stripped[1:].strip()
    stripped = re.sub(r"^\d+\)\s*", "", stripped)
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(tokens["bullet_space_after"])
    p.add_run("• ")
    if ": " in stripped:
        label, _, body = stripped.partition(":")
        if 1 <= len(label.split()) <= 5:
            label_run = p.add_run(label.strip() + ": ")
            label_run.bold = True
            _add_runs_with_metric_bold(p, body.strip())
            return
    _add_runs_with_metric_bold(p, stripped)


def _add_skill_bank_row(doc, line: str, tokens: dict | None = None):
    """Non-bulleted Skills keyword row with bold category label."""
    tokens = tokens or _doc_style_tokens()
    stripped = line.strip()
    if stripped.startswith(("-", "•")):
        stripped = stripped[1:].strip()
    stripped = re.sub(r"^\d+\)\s*", "", stripped)
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(max(2, tokens["bullet_space_after"] - 2))
    if ": " in stripped or (":" in stripped and not stripped.startswith("http")):
        label, _, body = stripped.partition(":")
        if 1 <= len(label.split()) <= 5:
            label_run = p.add_run(label.strip() + ": ")
            label_run.bold = True
            p.add_run(body.strip())
            return
    p.add_run(stripped)


def _add_role_header(doc, title_company: str, dates_location: str = "", tokens: dict | None = None):
    """Title bold left, dates 9pt right — matches PDF role-header."""
    tokens = tokens or _doc_style_tokens()
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(tokens["role_space_before"])
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.4), WD_TAB_ALIGNMENT.RIGHT)
    run_title = p.add_run(title_company)
    run_title.bold = True
    if dates_location:
        p.add_run("\t")
        run_dates = p.add_run(dates_location)
        run_dates.font.size = Pt(9)


def _add_company_desc_para(doc, text: str, tokens: dict | None = None):
    tokens = tokens or _doc_style_tokens()
    plain = _strip_html_tags(text)
    words = plain.split()
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(tokens["company_desc_space_after"])
    n = min(4, len(words))
    if not words:
        return
    if len(words) <= n:
        run = p.add_run(plain)
        run.bold = True
        run.italic = True
        run.font.size = Pt(9)
        return
    head = " ".join(words[:n])
    tail = " ".join(words[n:])
    run_head = p.add_run(head)
    run_head.bold = True
    run_head.italic = True
    run_head.font.size = Pt(9)
    run_tail = p.add_run(" " + tail)
    run_tail.italic = True
    run_tail.font.size = Pt(9)


def _add_role_block_doc(doc, role: dict, tokens: dict | None = None):
    tokens = tokens or _doc_style_tokens()
    if role.get("page_break_before"):
        _add_para(doc, "", style="Normal", page_break_before=True)

    if role.get("is_condensed"):
        _add_role_header(doc, role["title_company"], role.get("dates_location", ""), tokens=tokens)
        return

    _add_role_header(doc, role["title_company"], role.get("dates_location", ""), tokens=tokens)
    if role.get("company_desc"):
        _add_company_desc_para(doc, role["company_desc"], tokens=tokens)
    if role.get("summary"):
        _add_para(doc, role["summary"], style="Normal", space_after_pt=tokens["role_summary_space_after"])
    for bullet_html in role.get("bullets") or []:
        _add_list_bullet(doc, _strip_html_tags(bullet_html), tokens=tokens)


def _parse_skills_lines(skills_raw: str) -> list[str]:
    lines_out = []
    for line in strip_coaching_notes(skills_raw).split("\n"):
        stripped = line.strip()
        if not stripped or _is_page_break_line(stripped):
            continue
        stripped = re.sub(r"^\d+\)\s*", "", stripped)
        stripped = re.sub(r"^-\s*", "", stripped)
        lines_out.append(stripped)
    return lines_out


def _parse_skill_bank_lines(skill_bank_raw: str) -> list[str]:
    """Category rows for Skills (ATS keyword wall)."""
    return _parse_skills_lines(skill_bank_raw)


def _parse_experience_blocks(experience_list, break_role_index):
    """Same chunking as pdf_generator.render_resume_pdf."""
    parsed_roles = []
    for idx, role_text in enumerate(experience_list, start=1):
        role_break = break_role_index is not None and idx == break_role_index
        chunks = re.split(
            rf"(?m)^\s*{re.escape(PAGE_BREAK_MARKER)}\s*$",
            role_text,
        )
        for chunk_i, chunk in enumerate(chunks):
            chunk = chunk.strip()
            if not chunk:
                continue
            want_break = role_break or chunk_i > 0
            parsed = _parse_role_block(chunk, page_break_before=want_break)
            if parsed:
                if not parsed_roles and parsed.get("page_break_before"):
                    parsed = {**parsed, "page_break_before": False}
                parsed_roles.append(parsed)
    return parsed_roles


def _parse_projects_blocks(projects_raw: str) -> list[dict]:
    projects_list = []
    if not projects_raw:
        return projects_list
    proj_clean = strip_coaching_notes(projects_raw)
    current_title = None
    current_bullets = []
    pending_proj_break = False
    for line in proj_clean.split("\n"):
        stripped = line.strip()
        if _is_page_break_line(stripped):
            if current_title is not None:
                pending_proj_break = True
            continue
        if not stripped:
            if current_title:
                proj_break = pending_proj_break and bool(projects_list)
                projects_list.append({
                    "title": current_title,
                    "bullets": current_bullets,
                    "page_break_before": proj_break,
                })
                current_title = None
                current_bullets = []
                pending_proj_break = False
            continue
        if not stripped.startswith("-") and current_title is None:
            current_title = stripped
        else:
            bullet = re.sub(r"^-\s*", "", stripped)
            current_bullets.append(bullet)
    if current_title:
        proj_break = pending_proj_break and bool(projects_list)
        projects_list.append({
            "title": current_title,
            "bullets": current_bullets,
            "page_break_before": proj_break,
        })
    return projects_list


def create_formatted_doc(
    output_path,
    resume_sections,
    location="Sunnyvale, CA",
    include_scrum=False,
    pdf_breaks=None,
    export_mode="standard",
):
    """
    Build resume DOCX using assets/template.docx styles; structure matches template.html (PDF).
    """
    breaks = normalize_pdf_breaks(pdf_breaks)
    break_sections = breaks["before_sections"]
    break_role_index = breaks["before_role_index"]
    tokens = _doc_style_tokens(export_mode)

    doc = Document(str(TEMPLATE_PATH))
    _clear_body(doc)

    # Header — matches template.html .header
    p_name = doc.add_paragraph(style="Normal")
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run('Bharat "Bob" Lavania')
    run_name.font.name = NAME_FONT
    run_name.font.size = Pt(tokens["name_size"])
    run_name.bold = True

    mission_text = normalize_quick_take_text(resume_sections.get("mission", ""))
    mission_lines = [l.strip() for l in mission_text.split("\n") if l.strip()]
    tagline = mission_lines[0] if mission_lines else ""
    if tagline and " - " in tagline and ": " not in tagline[: tagline.find(" - ") + 1]:
        tagline = tagline.replace(" - ", ": ", 1)
    mission_body = "\n".join(mission_lines[1:]) if len(mission_lines) > 1 else ""

    if tagline:
        p_tag = doc.add_paragraph(style="Normal")
        p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tag = p_tag.add_run(tagline)
        run_tag.bold = True
        run_tag.font.size = Pt(tokens["tagline_size"])

    p_contact = doc.add_paragraph(style="Normal")
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_contact = p_contact.add_run(
        f"(312) 772-7962 | xblavania@gmail.com | linkedin.com/in/blavania | {location}"
    )
    run_contact.font.size = Pt(tokens["contact_size"])

    # Summary = positioning paragraph + former How I Work method bullets (bolstered, not replaced)
    skills_lines = _parse_skills_lines(resume_sections.get("skills", ""))
    if mission_body or skills_lines:
        _add_section_heading(doc, SECTION_SUMMARY, tokens=tokens)
        if mission_body:
            _add_para(doc, mission_body, style="Normal", space_after_pt=tokens["body_space_after"])
        for skill_line in skills_lines:
            _add_list_bullet(doc, skill_line, tokens=tokens)

    # Skills (ATS keyword wall)
    skill_bank_lines = _parse_skill_bank_lines(resume_sections.get("skill_bank", ""))
    if skill_bank_lines:
        _add_section_heading(
            doc,
            SECTION_SKILLS,
            page_break_before="skill_bank" in break_sections,
            tokens=tokens,
        )
        for row in skill_bank_lines:
            _add_skill_bank_row(doc, row, tokens=tokens)

    # Experience
    experience_raw = resume_sections.get("experience") or []
    if experience_raw:
        exp_page_break = "experience" in break_sections
        _add_section_heading(doc, SECTION_EXPERIENCE, page_break_before=exp_page_break, tokens=tokens)
        for role in _parse_experience_blocks(experience_raw, break_role_index):
            _add_role_block_doc(doc, role, tokens=tokens)

    # School Projects
    projects_raw = resume_sections.get("projects", "")
    if projects_raw and strip_coaching_notes(projects_raw).strip():
        _add_section_heading(
            doc,
            SECTION_SCHOOL_PROJECTS,
            page_break_before="projects" in break_sections,
            tokens=tokens,
        )
        for proj in _parse_projects_blocks(projects_raw):
            if proj.get("page_break_before"):
                _add_para(doc, "", style="Normal", page_break_before=True)
            _add_role_header(doc, proj["title"], "", tokens=tokens)
            for bullet in proj["bullets"]:
                _add_list_bullet(doc, bullet, tokens=tokens)

    # Credentials
    _add_section_heading(
        doc,
        SECTION_CREDENTIALS,
        page_break_before="credentials" in break_sections,
        tokens=tokens,
    )
    p_mba = doc.add_paragraph(style="Normal")
    r1 = p_mba.add_run("MBA, Energy Economics and Product Leadership")
    r1.bold = True
    p_mba.add_run(", University of Calgary (2025)")

    p_bs = doc.add_paragraph(style="Normal")
    r2 = p_bs.add_run("Bachelor of Science, Finance")
    r2.bold = True
    p_bs.add_run(", Illinois Tech, Chicago, IL (2012)")

    if include_scrum:
        p_scrum = doc.add_paragraph(style="Normal")
        r3 = p_scrum.add_run("CSM & CSPO")
        r3.bold = True
        p_scrum.add_run(", Scrum Alliance (2022)")

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path
